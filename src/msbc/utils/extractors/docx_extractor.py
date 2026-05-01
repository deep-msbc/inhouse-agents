"""
DOCX text and heading extraction using python-docx.

Public API
----------
extract_text(file_bytes)              -> str
    Full plain text with Markdown-style structural markers.

extract_heading_hierarchy(file_bytes) -> list[dict]
    Ordered list of {"level": int, "text": str} from Heading styles.
    Level 0 = Title style.  Level 1-6 = Heading 1 – Heading 6.
    Falls back to heuristic extraction when no Heading styles are found.
"""

import io
import logging
import re

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
# Public functions
# ─────────────────────────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes) -> str:
    """
    Extract plain text from a .docx file preserving document structure.

    - Heading styles   → Markdown markers  (# / ## / ### …)
    - Title style      → # Title
    - List paragraphs  → "  - item"
    - Tables           → Markdown table format
    - Paragraphs and tables are interleaved in document order.

    Returns:
        Cleaned, structured plain text string.

    Raises:
        RuntimeError: If python-docx is not installed.
    """
    doc = _open_docx(file_bytes)

    para_map  = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    parts: list[str] = []

    for child in doc.element.body:
        if child in para_map:
            para = para_map[child]
            text = para.text.strip()
            if not text:
                continue

            style_name  = (para.style.name or "") if para.style else ""
            lower_style = style_name.lower()

            if lower_style == "title":
                parts.append(f"\n# {text}")
            elif "heading" in lower_style:
                level = _heading_level(style_name)
                parts.append(f"\n{'#' * level} {text}")
            elif "list" in lower_style or _has_numbering(para):
                parts.append(f"  - {text}")
            else:
                parts.append(text)

        elif child in table_map:
            _format_table(table_map[child], parts)

    return _clean_text("\n".join(parts))


def extract_heading_hierarchy(file_bytes: bytes) -> list[dict]:
    """
    Extract document headings using a dual-layer approach.

    Layer 1 (style-based):
        Paragraphs with Word paragraph styles "Title" or "Heading 1–6".
        Most reliable when the author applied proper styles.

    Layer 2 (format-based):
        Paragraphs detected as headings purely from visual formatting:
          a. XML outline level (w:outlineLvl) — set by Word's outline tool
             even when a custom/unnamed style is used.
          b. All-bold runs + numbered section pattern (e.g. "8. Material Issue",
             "3.2.4 Hardware Tab") — the most common pattern in ERP user stories
             where authors type bold numbered headings without applying styles.
          c. Large font (≥ 14 pt) + all-bold + short text — catch visually
             prominent section titles that don't fit the numbered pattern.

    Both layers always run. Results are merged in document order and
    deduplicated by text (style-detected level wins when both catch the
    same paragraph). Falls back to heuristic extraction from plain text
    only when both layers return nothing.

    Returns:
        Ordered list of {"level": int, "text": str} dicts.
        level 0 = Title, level 1-6 = Heading 1 through Heading 6.

    Raises:
        RuntimeError: If python-docx is not installed.
    """
    doc = _open_docx(file_bytes)

    style_entries:  list[tuple[int, dict]] = []   # (para_idx, heading_dict)
    format_entries: list[tuple[int, dict]] = []
    style_texts:    dict[str, int]         = {}   # text → level (for dedup)

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        style_name  = (para.style.name or "") if para.style else ""
        lower_style = style_name.lower()

        # ── Layer 1: Style-based ──────────────────────────────────────────────
        if lower_style == "title":
            entry = {"level": 0, "text": text}
            style_entries.append((para_idx, entry))
            style_texts[text] = 0
            continue

        if "heading" in lower_style:
            level = _heading_level(style_name)
            entry = {"level": level, "text": text}
            style_entries.append((para_idx, entry))
            style_texts[text] = level
            continue

        # ── Layer 2: Format-based (only for paragraphs not caught by Layer 1) ─
        if text in style_texts:
            continue

        detected_level = _detect_format_heading_level(para, text)
        if detected_level is not None:
            format_entries.append((para_idx, {"level": detected_level, "text": text}))

    # ── Merge in document order ───────────────────────────────────────────────
    all_entries = style_entries + format_entries
    all_entries.sort(key=lambda x: x[0])
    headings = [entry for _, entry in all_entries]

    if not headings:
        logger.warning(
            "extract_heading_hierarchy: no headings found via style or format "
            "detection. Falling back to heuristic extraction from extracted text."
        )
        return _heuristic_headings(extract_text(file_bytes))

    style_count  = len(style_entries)
    format_count = len(format_entries)
    if format_count:
        logger.info(
            "extract_heading_hierarchy: %d style-based + %d format-based headings"
            " (%d total).",
            style_count, format_count, len(headings),
        )

    return headings


# ─────────────────────────────────────────────────────────────────────────────
# Internal helpers
# ─────────────────────────────────────────────────────────────────────────────

def _open_docx(file_bytes: bytes):
    """Open a DOCX from bytes. Raises RuntimeError if python-docx is missing."""
    try:
        from docx import Document  # python-docx
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )
    return Document(io.BytesIO(file_bytes))


def _heading_level(style_name: str) -> int:
    """
    Extract numeric level from a style name like 'Heading 2' → 2.
    Clamps to [1, 6]. Defaults to 1 if no digit found.
    """
    match = re.search(r"\d+", style_name)
    return min(int(match.group()), 6) if match else 1


def _has_numbering(para) -> bool:
    """Return True if the paragraph carries list-numbering XML (numPr element)."""
    pPr = para._element.pPr
    if pPr is None:
        return False
    return any(child.tag.endswith("}numPr") for child in pPr)


def _format_table(table, parts: list[str]) -> None:
    """
    Render a DOCX table as a Markdown-style table and append to parts.

    Horizontally merged cells (same XML element appearing multiple times in
    a row) are deduplicated to avoid repeated text in the same row.
    """
    parts.append("")  # blank line before table
    for row_idx, row in enumerate(table.rows):
        seen_ids: set[int] = set()
        unique_cells: list[str] = []
        for cell in row.cells:
            cell_id = id(cell._element)
            if cell_id not in seen_ids:
                seen_ids.add(cell_id)
                unique_cells.append(cell.text.strip().replace("\n", " "))
        parts.append("| " + " | ".join(unique_cells) + " |")
        if row_idx == 0:
            parts.append("|" + " --- |" * len(unique_cells))
    parts.append("")  # blank line after table


def _clean_text(text: str) -> str:
    """
    Normalize whitespace without removing any meaningful content.

    Steps:
    1. Strip null bytes and ASCII control chars (except tab and newline).
    2. Collapse runs of horizontal whitespace to a single space per line.
    3. Strip trailing whitespace from every line.
    4. Collapse 3+ consecutive blank lines down to 2.
    """
    if not text:
        return text

    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"[^\S\n]+", " ", text)
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def _detect_format_heading_level(para, text: str) -> int | None:
    """
    Infer heading level from paragraph formatting (no reliance on style names).

    Checks (in priority order):
    1. XML outline level (w:outlineLvl) — authoritative even with custom styles.
    2. Numbered section pattern + all-bold runs — e.g. "8. Material Issue".
    3. Large font (≥ 14 pt) + all-bold + short text — visual prominence cue.

    Returns a level in [1, 6] or None if the paragraph is not a heading.
    """
    # 1. Outline level
    outline_level = _para_outline_level(para)
    if outline_level is not None:
        # OOXML: 0 → Heading 1, 1 → Heading 2, … 5 → Heading 6
        return min(outline_level + 1, 6)

    # 2. Numbered section + bold
    numbered_level = _numbered_section_level(text)
    if numbered_level is not None and _is_bold_paragraph(para):
        return numbered_level

    # 3. Large font + bold + short
    font_pt = _para_font_size_pt(para)
    if (
        font_pt is not None
        and font_pt >= 14
        and _is_bold_paragraph(para)
        and len(text.split()) <= 12
        and len(text) <= 100
    ):
        return 2  # treat as top-level heading

    return None


def _para_outline_level(para) -> int | None:
    """Return the OOXML w:outlineLvl value (0–8) or None if not set."""
    try:
        pPr = para._element.pPr
        if pPr is None:
            return None
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        outlineLvl = pPr.find(f"{{{ns}}}outlineLvl")
        if outlineLvl is None:
            return None
        val = outlineLvl.get(f"{{{ns}}}val")
        if val is None:
            return None
        level = int(val)
        return level if level <= 8 else None
    except Exception:
        return None


def _is_bold_paragraph(para) -> bool:
    """Return True if all non-empty runs in the paragraph are bold."""
    runs = [r for r in para.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def _para_font_size_pt(para) -> float | None:
    """Return font size in points from the paragraph's first non-empty run."""
    for run in para.runs:
        if run.text.strip() and run.font.size:
            try:
                return run.font.size.pt
            except Exception:
                pass
    return None


def _numbered_section_level(text: str) -> int | None:
    """
    Detect a numbered-section heading and return its depth level.

    Examples:
        "8. Material Issue With Job"    → 2  (1 numeric part → level 2)
        "8.1 Sub-section"               → 3  (2 parts → level 3)
        "3.2.4 Hardware Tab"            → 5  (4 parts → level 5)
    Returns None if *text* does not match the numbered-section pattern.
    """
    m = re.match(r"^(\d+(?:\.\d+)*)\.?\s+\S", text)
    if not m:
        return None
    depth = len(m.group(1).split("."))
    return min(depth + 1, 6)


def _heuristic_headings(text: str) -> list[dict]:
    """
    Fallback heading detector for DOCX files that lack proper Heading styles.

    Detects (in priority order):
      1. Markdown headings  (#, ##, ###, ####)
      2. Numbered sections  (1. Title / 1.1 Title / 1.1.1. Title)
      3. RST-style underlined titles  (next line is === or ---)
      4. ALL-CAPS lines (≥10 chars)
      5. Short standalone lines followed by a blank line

    Returns:
        [{level: int, text: str}] in document order.
    """
    lines = text.splitlines()
    headings: list[dict] = []
    n = len(lines)
    skip_next = False

    for i, line in enumerate(lines):
        if skip_next:
            skip_next = False
            continue

        stripped = line.strip()
        if not stripped:
            continue

        # 1. Markdown headings
        m = re.match(r"^(#{1,6})\s+(.+)", stripped)
        if m:
            headings.append({"level": len(m.group(1)), "text": m.group(2).strip()})
            continue

        # 2. Numbered sections
        if re.match(r"^\d+(\.|\d+)*\.?\s+[A-Z\u0080-\uFFFF]", stripped) or re.match(
            r"^(Chapter|Section|Part|Module)\s+\d+", stripped, re.IGNORECASE
        ):
            headings.append({"level": 2, "text": stripped})
            continue

        # 3. RST underlined titles
        if i + 1 < n:
            next_s = lines[i + 1].strip()
            if (
                next_s
                and re.match(r"^[=\-]{3,}$", next_s)
                and len(next_s) >= max(3, len(stripped) // 2)
            ):
                headings.append({"level": 1, "text": stripped})
                skip_next = True
                continue

        # 4. ALL-CAPS line
        if (
            len(stripped) >= 10
            and stripped.upper() == stripped
            and re.search(r"[A-Z]{3,}", stripped)
            and not re.fullmatch(r"[^a-zA-Z0-9]*", stripped)
        ):
            headings.append({"level": 1, "text": stripped})
            continue

        # 5. Short standalone line followed by a blank line
        if (
            2 <= len(stripped.split()) <= 12
            and len(stripped) <= 80
            and stripped[-1] not in (".", ",", ";", ":")
            and not re.search(r"\w\.\s+\w", stripped)
            and i + 1 < n
            and not lines[i + 1].strip()
        ):
            headings.append({"level": 2, "text": stripped})

    return headings
